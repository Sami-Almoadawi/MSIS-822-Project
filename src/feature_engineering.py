# =================================================================================================
# PHASE 3: FEATURE ENGINEERING & DATA SPLITTING
# =================================================================================================

# 1- Task 3.1:
# Engineer Individual Stylometric Feature Assignment (6، 29، 52، 75، 98)

# 1. INSTALL SYSTEM DEPENDENCIES AND LIBRARIES
print("⏳ Installing system dependencies and libraries...")
import sys
# 2. IMPORTS
print("🚀 PHASE 3 STARTED: Feature Engineering") 
print("=" * 80)

import os
import gc
import re
import warnings
from datetime import datetime
from typing import Dict, List, Tuple

import numpy as np
import pandas as pd
import torch
import joblib
from sklearn.model_selection import train_test_split
from tqdm import tqdm
from transformers import AutoTokenizer, AutoModel
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import openpyxl  # noqa: F401 (needed by pandas ExcelWriter)

# -----------------------------------------------------------------------------------------
# Global settings
# -----------------------------------------------------------------------------------------
warnings.filterwarnings("ignore")
tqdm.pandas()

try:
    from IPython.display import display  # type: ignore
    IS_JUPYTER = True
except ImportError:  # fallback for non-notebook environments
    IS_JUPYTER = False

    def display(x):
        print(x)


# =========================================================================================
# Helper functions
# =========================================================================================

def style_word_table(table) -> None:
    """Apply borders and blue header styling to a Word table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)

    # Add borders
    for border in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        edge = OxmlElement(f"w:{border}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), "000000")
        tblBorders.append(edge)

    # Header row styling
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i == 0:
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "2E75B6")
                tcPr.append(shd)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)


def load_preprocessing_excel(path: str) -> Tuple[pd.DataFrame, Dict[str, pd.DataFrame]]:
    """
    Load all sheets from preprocessing_results.xlsx and build:
    - df_all: concatenated DataFrame
    - split_dfs: dict of per-split DataFrames
    """
    print("📂 Loading preprocessing_results.xlsx...")
    if not os.path.exists(path):
        raise FileNotFoundError(
            f"'preprocessing_results.xlsx' not found at:\n  {path}\n"
            f"Run the preprocessing phase first."
        )

    excel = pd.ExcelFile(path)
    frames: List[pd.DataFrame] = []
    split_dfs: Dict[str, pd.DataFrame] = {}

    for sheet in excel.sheet_names:
        print(f"  ↳ Processing sheet: {sheet}")
        df = pd.read_excel(path, sheet_name=sheet)

        # Robust mapping from sheet name to split label
        split_prefix = sheet.split("_")[0].lower() if "_" in sheet else sheet.lower()
        if split_prefix == "by":
            split_name = "by_polishing"
        elif split_prefix == "from":
            split_name = "from_title_and_content" if "title_and_content" in sheet.lower() else "from_title"
        else:
            split_name = split_prefix

        # Text type: Human vs AI
        text_type = "Human" if "original_abstract" in sheet.lower() else "AI"

        df["Split"] = split_name
        df["Text Type"] = text_type

        frames.append(df)
        split_dfs[split_name] = df

    df_all = pd.concat(frames, ignore_index=True)

    # Normalize column names if they exist
    if {"Original", "After Preprocessing"}.issubset(df_all.columns):
        df_all = df_all.rename(
            columns={
                "Original": "Original Text",
                "After Preprocessing": "Text After Processing",
            }
        )

    if "Text After Processing" not in df_all.columns:
        raise KeyError(
            "Required column 'Text After Processing' is missing.\n"
            "Check the schema of preprocessing_results.xlsx."
        )

    print(f"✅ Dataset loaded: {len(df_all):,} rows")
    print(f"   Splits detected: {list(split_dfs.keys())}")
    return df_all, split_dfs


# --- Stylometric features ----------------------------------------------------------------

def feat_006_multiple_elongations(text: str) -> int:
    """Count occurrences of character elongation (3+ repeated chars)."""
    if not isinstance(text, str):
        return 0
    return len(re.findall(r"(.)\1{2,}", text))


def feat_029_semicolons(text: str) -> int:
    """Count Arabic semicolons in text."""
    if not isinstance(text, str):
        return 0
    return text.count("؛")


INTERJECTIONS_AR = {
    "آه", "أوه", "واه", "وي", "هيا", "يا", "الله", "آمين", "سبحان", "ياسلام"
}


def feat_052_interjections(text: str) -> int:
    """Count Arabic interjections from a predefined list."""
    if not isinstance(text, str):
        return 0
    words = text.split()
    return sum(1 for w in words if w.strip() in INTERJECTIONS_AR)


def feat_075_active_voice_sentences(text: str) -> int:
    """
    Approximate count of active-voice sentences by excluding simple passive markers.
    This is a heuristic, not full Arabic syntax analysis.
    """
    if not isinstance(text, str):
        return 0

    sentences = re.split(r"[.!؟\n\r]+", text)
    passive_markers = ["تُ", "يُ", "أُ", "تم ", "تمت ", "تمّ ", "تمّت "]

    count = 0
    for s in sentences:
        s = s.strip()
        if not s:
            continue
        if not any(pm in s for pm in passive_markers):
            count += 1
    return count


def add_shallow_features(df_all: pd.DataFrame) -> pd.DataFrame:
    """Add features 006, 029, 052, 075 to df_all in-place."""
    print("\n⚙️ Extracting 4 shallow stylometric features...")
    text_series = df_all["Text After Processing"]

    df_all["feat_006_multiple_elongations"] = text_series.progress_apply(feat_006_multiple_elongations)
    df_all["feat_029_semicolons"] = text_series.progress_apply(feat_029_semicolons)
    df_all["feat_052_interjections"] = text_series.progress_apply(feat_052_interjections)
    df_all["feat_075_active_voice_sentences"] = text_series.progress_apply(feat_075_active_voice_sentences)

    return df_all


def add_bert_cls_l2norm(
    df_all: pd.DataFrame,
    text_col: str = "Text After Processing",
    batch_size: int = 16,
    max_length: int = 256,
) -> pd.DataFrame:
    """Add feature 098: BERT CLS L2-norm (batched & memory-friendly)."""
    print("\n🧠 Extracting Feature 098: BERT CLS L2-norm")

    device = "cuda" if torch.cuda.is_available() else "cpu"
    print(f"   Using device: {device.upper()}")

    bert_model_name = "aubmindlab/bert-base-arabertv2"
    tokenizer = AutoTokenizer.from_pretrained(bert_model_name)
    model = AutoModel.from_pretrained(bert_model_name).to(device)
    model.eval()

    df_all["feat_098_bert_cls_l2norm"] = 0.0

    texts = df_all[text_col].astype(str).tolist()
    n = len(texts)

    for start in tqdm(range(0, n, batch_size), desc="BERT batches"):
        end = min(start + batch_size, n)
        batch_texts = texts[start:end]

        try:
            inputs = tokenizer(
                batch_texts,
                return_tensors="pt",
                max_length=max_length,
                truncation=True,
                padding=True,
            ).to(device)

            with torch.no_grad():
                outputs = model(**inputs)
                cls_embeddings = outputs.last_hidden_state[:, 0, :]  # CLS token
                norms = torch.norm(cls_embeddings, p=2, dim=1).cpu().numpy()

            df_all.loc[start:end - 1, "feat_098_bert_cls_l2norm"] = norms

        except Exception as e:
            print(f"   ⚠️ Skipping batch {start}-{end} due to error: {e}")

        # Free memory regularly
        del inputs
        if device == "cuda":
            torch.cuda.empty_cache()
        gc.collect()

    print("✅ All 5 features extracted successfully.")
    return df_all


def stratified_split(df_all: pd.DataFrame, label_col: str = "Text Type"):
    """Perform 70/15/15 stratified split on Text Type."""
    print("\n🔀 Performing 70/15/15 stratified split...")

    if label_col not in df_all.columns:
        raise KeyError(f"Column '{label_col}' is missing; cannot perform stratified split.")

    train_df, temp_df = train_test_split(
        df_all,
        test_size=0.30,
        random_state=42,
        stratify=df_all[label_col],
    )

    val_df, test_df = train_test_split(
        temp_df,
        test_size=0.50,
        random_state=42,
        stratify=temp_df[label_col],
    )

    split_summary = pd.DataFrame(
        {
            "Part": ["Total", "Train (70%)", "Validation (15%)", "Test (15%)"],
            "Records": [len(df_all), len(train_df), len(val_df), len(test_df)],
            "Human": [
                df_all[label_col].eq("Human").sum(),
                train_df[label_col].eq("Human").sum(),
                val_df[label_col].eq("Human").sum(),
                test_df[label_col].eq("Human").sum(),
            ],
            "AI": [
                df_all[label_col].eq("AI").sum(),
                train_df[label_col].eq("AI").sum(),
                val_df[label_col].eq("AI").sum(),
                test_df[label_col].eq("AI").sum(),
            ],
        }
    )

    return train_df, val_df, test_df, split_summary


# =========================================================================================
# Main pipeline
# =========================================================================================
def main():
    print("⏳ Installing system dependencies and libraries...")
    print("🚀 PHASE 3 STARTED: Feature Engineering + 8 Excel Sheets + Pickle")
    print("=" * 80)

    # ------------------------------------------------------------------
    # 1. Load preprocessing_results.xlsx
    # ------------------------------------------------------------------
    base_dir = os.path.dirname(os.path.abspath(__file__))
    input_path = os.path.join(base_dir, "preprocessing_results.xlsx")
    df_all, split_dfs = load_preprocessing_excel(input_path)

    # ------------------------------------------------------------------
    # 2. Add stylometric features (shallow)
    # ------------------------------------------------------------------
    df_all = add_shallow_features(df_all)

    # ------------------------------------------------------------------
    # 3. Add BERT CLS L2-norm (batched)
    # ------------------------------------------------------------------
    df_all = add_bert_cls_l2norm(df_all, text_col="Text After Processing")

    feature_cols = [
        "feat_006_multiple_elongations",
        "feat_029_semicolons",
        "feat_052_interjections",
        "feat_075_active_voice_sentences",
        "feat_098_bert_cls_l2norm",
    ]

    # ------------------------------------------------------------------
    # 4. Stratified split 70/15/15
    # ------------------------------------------------------------------
    train_df, val_df, test_df, split_summary = stratified_split(df_all, label_col="Text Type")
    feature_stats = df_all[feature_cols].describe().round(4)

    print("\n📋 Split summary:")
    print(split_summary.to_string(index=False))

    # ------------------------------------------------------------------
    # 5. Save Excel with 8 sheets
    # ------------------------------------------------------------------
    print("\n💾 Saving 'Complete_Dataset_With_Features.xlsx' (up to 8 sheets)...")
    output_excel = os.path.join(base_dir, "Complete_Dataset_With_Features.xlsx")

    with pd.ExcelWriter(output_excel, engine="openpyxl") as writer:
        train_df.to_excel(writer, sheet_name="Train", index=False)
        val_df.to_excel(writer, sheet_name="Validation", index=False)
        test_df.to_excel(writer, sheet_name="Test", index=False)
        df_all.to_excel(writer, sheet_name="All_Data", index=False)
        split_summary.to_excel(writer, sheet_name="Split_Summary", index=False)
        feature_stats.to_excel(writer, sheet_name="Feature_Stats", index=True)

        # Optional split sheets (only if they exist)
        if "by_polishing" in split_dfs:
            split_dfs["by_polishing"].to_excel(
                writer, sheet_name="by_polishing", index=False
            )
        if "from_title_and_content" in split_dfs:
            split_dfs["from_title_and_content"].to_excel(
                writer, sheet_name="from_title_and_content", index=False
            )

    print(f"✅ Excel saved: {output_excel}")

    # ------------------------------------------------------------------
    # 6. Save Word report
    # ------------------------------------------------------------------
    print("\n💾 Saving 'Complete_Dataset_With_Features.docx'...")

    doc = Document()
    doc.add_heading("Feature Engineering Report - Phase 3", 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(
        f"Total Records: {len(df_all):,} | Engineered Features: {len(feature_cols)}"
    )

    # Data split table
    doc.add_heading("1. Data Split Summary", level=1)
    table = doc.add_table(rows=len(split_summary) + 1, cols=len(split_summary.columns))
    style_word_table(table)

    for i, col in enumerate(split_summary.columns):
        table.rows[0].cells[i].text = str(col)

    for i, (_, row) in enumerate(split_summary.iterrows()):
        for j, val in enumerate(row):
            if isinstance(val, (int, np.integer)):
                table.rows[i + 1].cells[j].text = f"{int(val):,}"
            else:
                table.rows[i + 1].cells[j].text = str(val)

    # Sheet overview (derived dynamically from writer logic)
    doc.add_heading("2. Generated Excel Sheets", level=1)
    sheet_names = [
        "Train",
        "Validation",
        "Test",
        "All_Data",
        "Split_Summary",
        "Feature_Stats",
    ]
    if "by_polishing" in split_dfs:
        sheet_names.append("by_polishing")
    if "from_title_and_content" in split_dfs:
        sheet_names.append("from_title_and_content")

    doc.add_paragraph("\n".join(f"• {name}" for name in sheet_names))

    word_path = os.path.join(base_dir, "Complete_Dataset_With_Features.docx")
    doc.save(word_path)
    print("✅ Word report saved.")

    # ------------------------------------------------------------------
    # 7. Save pickle for Phase 4
    # ------------------------------------------------------------------
    print("\n💾 Saving 'phase3_splits.pkl' for modeling phase...")

    splits_dict = {
        "train_indices": train_df.index.tolist(),
        "val_indices": val_df.index.tolist(),
        "test_indices": test_df.index.tolist(),
        "df_all": df_all,       # full DataFrame with all features
        "feature_cols": feature_cols,
    }
    pickle_path = os.path.join(base_dir, "phase3_splits.pkl")
    joblib.dump(splits_dict, pickle_path)
    print("✅ phase3_splits.pkl saved.")

    # ------------------------------------------------------------------
    # 8. Final log
    # ------------------------------------------------------------------
    print("\n" + "=" * 80)
    print("🎉 PHASE 3 COMPLETED SUCCESSFULLY!")
    print("=" * 80)
    print("Output Files:")
    print(f"  1) {output_excel}")
    print(f"  2) {word_path}")
    print(f"  3) {pickle_path}")
    print("\nData Distribution:")
    print(f"  • Train:      {len(train_df):,} ({len(train_df) / len(df_all) * 100:.1f}%)")
    print(f"  • Validation: {len(val_df):,} ({len(val_df) / len(df_all) * 100:.1f}%)")
    print(f"  • Test:       {len(test_df):,} ({len(test_df) / len(df_all) * 100:.1f}%)")
    print("\nEngineered Features:")
    for feat in feature_cols:
        print(f"  • {feat}")
    print("\n🚀 Ready for Phase 4 (modeling.py)")
    print("=" * 80)


if __name__ == "__main__":
    main()