# =================================================================================================
# PHASE 4: MACHINE LEARNING CLASSIFICATION MODELS
# =================================================================================================


# IMPORTS

import os
import zipfile
import warnings

import numpy as np
import pandas as pd
import joblib

import matplotlib.pyplot as plt
import seaborn as sns

from sklearn.model_selection import train_test_split
from sklearn.preprocessing import LabelEncoder, StandardScaler
from sklearn.linear_model import LogisticRegression
from sklearn.naive_bayes import GaussianNB
from sklearn.ensemble import RandomForestClassifier
from sklearn.svm import SVC
from xgboost import XGBClassifier
from sklearn.metrics import (
    accuracy_score,
    confusion_matrix,
    classification_report,
    roc_auc_score,
    roc_curve,
    auc,
    f1_score,
)

from tensorflow.keras.models import Sequential
from tensorflow.keras.layers import Dense, Dropout, BatchNormalization

from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

warnings.filterwarnings("ignore")
plt.style.use("default")
sns.set_palette("viridis")


# -----------------------------------------------------------------------------------------
# Helper: style Word tables
# -----------------------------------------------------------------------------------------
def style_word_table(table):
    """Apply professional styling to Word tables (borders + blue header)."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)

    for border in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        edge = OxmlElement(f"w:{border}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), "000000")
        tblBorders.append(edge)

    for i, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i == 0:
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "2E75B6")
                tcPr.append(shd)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)


# -----------------------------------------------------------------------------------------
# Helper: save models
# -----------------------------------------------------------------------------------------
def save_all_models(models_dict, save_dir="models"):
    """
    Save all ML/DL models in native formats:
    - .pkl for traditional ML models
    - .keras for Keras models
    Returns a dict {logical_name: relative_path}.
    """
    os.makedirs(save_dir, exist_ok=True)
    saved_paths = {}

    for name, model in models_dict.items():
        if "Sequential" in str(type(model)):
            path = os.path.join(save_dir, f"{name}.keras")
            model.save(path)
            print(f"  [Saved NN]  {path}")
        else:
            path = os.path.join(save_dir, f"{name}.pkl")
            joblib.dump(model, path)
            print(f"  [Saved ML]  {path}")
        saved_paths[name] = os.path.basename(path)

    print("✅ All models saved.")
    return saved_paths


# -----------------------------------------------------------------------------------------
# Helper: zip models folder
# -----------------------------------------------------------------------------------------
def create_models_zip(save_dir="models", zip_name="models.zip"):
    """Create a compressed ZIP archive of the models folder."""
    if os.path.exists(zip_name):
        os.remove(zip_name)

    with zipfile.ZipFile(zip_name, "w", zipfile.ZIP_DEFLATED) as zipf:
        for root, dirs, files in os.walk(save_dir):
            for file in files:
                full_path = os.path.join(root, file)
                arcname = os.path.relpath(full_path, ".")
                zipf.write(full_path, arcname)
    size_kb = os.path.getsize(zip_name) / 1024
    print(f"✅ {zip_name} created ({size_kb:.1f} KB)")
    return zip_name


# -----------------------------------------------------------------------------------------
# Helper: unified predictor file
# -----------------------------------------------------------------------------------------
def save_unified_predictor(scaler, label_encoder, model_paths, best_model_key, save_dir="models"):
    """
    Save:
    - scaler.pkl
    - label_encoder.pkl
    - unified_ai_detector.py  (single entry point for inference)
    model_paths: {logical_name: filename_inside_models_dir}
    best_model_key: key in model_paths for best model (e.g., 'random_forest' or 'feedforward_nn')
    """
    os.makedirs(save_dir, exist_ok=True)

    scaler_path = os.path.join(save_dir, "scaler.pkl")
    le_path = os.path.join(save_dir, "label_encoder.pkl")
    joblib.dump(scaler, scaler_path)
    joblib.dump(label_encoder, le_path)

    # Build Python code for unified predictor
    lines = [
        "import os",
        "import joblib",
        "import numpy as np",
        "import tensorflow as tf",
        "",
        "",
        "class UnifiedAIDetector:",
        "    def __init__(self, models_dir='models'):",
        "        self.models_dir = models_dir",
        "        self.scaler = joblib.load(os.path.join(models_dir, 'scaler.pkl'))",
        "        self.label_encoder = joblib.load(os.path.join(models_dir, 'label_encoder.pkl'))",
        "",
        "    def _load_model(self, fname):",
        "        path = os.path.join(self.models_dir, fname)",
        "        if fname.endswith('.keras'):",
        "            return tf.keras.models.load_model(path), 'keras'",
        "        else:",
        "            return joblib.load(path), 'sklearn'",
        "",
        "    def predict_best_proba(self, X_new):",
        "        '''Return AI probability using the best model only.'''",
        "        X_scaled = self.scaler.transform(X_new)",
        f"        model, mtype = self._load_model('{model_paths[best_model_key]}')",
        "        if mtype == 'keras':",
        "            probs = model.predict(X_scaled, verbose=0).flatten()",
        "        else:",
        "            probs = model.predict_proba(X_scaled)[:, 1]",
        "        return probs",
        "",
        "    def predict_best_label(self, X_new, threshold=0.5):",
        "        '''Return labels (0/1) using the best model only.'''",
        "        probs = self.predict_best_proba(X_new)",
        "        return (probs > threshold).astype(int)",
        "",
        "    def predict_all_proba(self, X_new):",
        "        '''Return probabilities from all models as a dict.'''",
        "        X_scaled = self.scaler.transform(X_new)",
        "        results = {}",
    ]

    for logical_name, fname in model_paths.items():
        key = logical_name
        lines.extend(
            [
                f"        model, mtype = self._load_model('{fname}')",
                "        if mtype == 'keras':",
                "            probs = model.predict(X_scaled, verbose=0).flatten()",
                "        else:",
                "            probs = model.predict_proba(X_scaled)[:, 1]",
                f"        results['{key}'] = probs",
            ]
        )

    lines.extend(
        [
            "        return results",
            "",
            "",
            "detector = UnifiedAIDetector()",
            "print('🚀 UnifiedAIDetector ready. Use detector.predict_best_proba(X)')",
        ]
    )

    predictor_path = os.path.join(save_dir, "unified_ai_detector.py")
    with open(predictor_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    print(f"✅ Unified predictor saved: {predictor_path}")
    return predictor_path


# =========================================================================================
# 1. LOAD DATASET
# =========================================================================================
print("🚀 PHASE 4 STARTED: Modeling & Evaluation")
print("=" * 80)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
data_path = os.path.join(BASE_DIR, "Complete_Dataset_With_Features.xlsx")
sheet_name = "All_Data"

print("📂 Loading dataset...")
if not os.path.exists(data_path):
    raise FileNotFoundError(
        f"Required file '{data_path}' not found. Run Phase 3 first."
    )

try:
    df_all = pd.read_excel(data_path, sheet_name=sheet_name, engine="openpyxl")
    print(f"✅ Loaded {len(df_all)} rows from sheet '{sheet_name}'")
except Exception:
    df_all = pd.read_excel(data_path, engine="openpyxl")
    print(f"⚠️ Sheet '{sheet_name}' not found. Loaded default sheet: {len(df_all)} rows")

text_col = "Text After Processing" if "Text After Processing" in df_all.columns else df_all.columns[0]
orig_text_col = "Original Text" if "Original Text" in df_all.columns else text_col
label_col = "Text Type" if "Text Type" in df_all.columns else "label"

print(f"📊 Dataset shape: {df_all.shape}")
print(f"📝 Text column: {text_col}")
print(f"🏷️ Label column: {label_col}")

# =========================================================================================
# 2. FEATURE PREPARATION
# =========================================================================================
print("\n🔧 Preparing features...")

# Normalize BERT feature name (if Phase 3 changed)
feat98_name = "feat_098_bert_cls_l2norm"
if "feat_098_bert_cls_norm" in df_all.columns:
    df_all.rename(columns={"feat_098_bert_cls_norm": feat98_name}, inplace=True)
elif "feat_098_bert_cls_prob" in df_all.columns:
    df_all.rename(columns={"feat_098_bert_cls_prob": feat98_name}, inplace=True)

required_features = [
    "feat_006_multiple_elongations",
    "feat_029_semicolons",
    "feat_052_interjections",
    "feat_075_active_voice_sentences",
    feat98_name,
]

missing = [c for c in required_features if c not in df_all.columns]
if missing:
    print("Available columns:", list(df_all.columns))
    raise ValueError(f"Missing required features: {missing}")

indices = df_all.index.values
X = df_all[required_features].fillna(0).values
y = df_all[label_col].values

# Encode labels
le = LabelEncoder()
y_enc = le.fit_transform(y)

# Dynamic human/AI label mapping
label_map = {i: label for i, label in enumerate(le.classes_)}
ai_label_name = next((v for v in label_map.values() if "AI" in v.upper()), le.classes_[1])
human_label_name = next((v for v in label_map.values() if v != ai_label_name), le.classes_[0])
ai_label = le.transform([ai_label_name])[0]

print(f"ℹ️ Classes: Human='{human_label_name}', AI='{ai_label_name}'")
print(f"📈 Feature matrix: {X.shape}")

# Scaling
scaler = StandardScaler()
X_scaled = scaler.fit_transform(X)

# =========================================================================================
# 3. TRAIN / VALIDATION / TEST SPLIT (70 / 15 / 15)
# =========================================================================================
print("\n🎯 Creating 70/15/15 stratified split...")

idx_train, idx_temp, y_train, y_temp = train_test_split(
    indices, y_enc, test_size=0.30, random_state=42, stratify=y_enc
)
idx_val, idx_test, y_val, y_test = train_test_split(
    idx_temp, y_temp, test_size=0.50, random_state=42, stratify=y_temp
)

X_train_s, X_temp_s, _, _ = train_test_split(
    X_scaled, y_enc, test_size=0.30, random_state=42, stratify=y_enc
)
X_val_s, X_test_s, _, _ = train_test_split(
    X_temp_s, y_temp, test_size=0.50, random_state=42, stratify=y_temp
)

split_summary_df = pd.DataFrame(
    {
        "Subset": ["Training", "Validation", "Testing"],
        "Count": [len(idx_train), len(idx_val), len(idx_test)],
        "Ratio": ["70%", "15%", "15%"],
    }
)
print("\n📋 Data split summary:")
print(split_summary_df.to_string(index=False))

# =========================================================================================
# 4. TRAIN TRADITIONAL ML MODELS (batch 1)
# =========================================================================================
print("\n⚙️ Training traditional ML models...")

models = {
    "naive_bayes": GaussianNB(),
    "logistic_regression": LogisticRegression(max_iter=1000, random_state=42),
    "random_forest": RandomForestClassifier(n_estimators=100, random_state=42),
    "svm": SVC(kernel="rbf", probability=True, random_state=42),
    "xgboost": XGBClassifier(eval_metric="logloss", random_state=42),
}

results_list = []
model_preds = {}
model_probs = {}

for key, model in models.items():
    print(f"  {key}...", end=" ")
    model.fit(X_train_s, y_train)
    preds = model.predict(X_test_s)
    probs = model.predict_proba(X_test_s)[:, 1]

    model_preds[key] = preds
    model_probs[key] = probs

    report_dict = classification_report(
        y_test,
        preds,
        output_dict=True,
        target_names=le.classes_,
        zero_division=0,
    )
    ai_metrics = report_dict.get(ai_label_name, {})
    human_metrics = report_dict.get(human_label_name, {})

    results_list.append(
        {
            "Model": key.replace("_", " ").title(),
            "Accuracy": accuracy_score(y_test, preds),
            "Macro F1-Score": f1_score(y_test, preds, average="macro"),
            "ROC-AUC": roc_auc_score(y_test, probs),
            "Precision (AI)": ai_metrics.get("precision", 0),
            "Recall (AI)": ai_metrics.get("recall", 0),
            "F1-Score (AI)": ai_metrics.get("f1-score", 0),
            "Precision (Human)": human_metrics.get("precision", 0),
            "Recall (Human)": human_metrics.get("recall", 0),
            "F1-Score (Human)": human_metrics.get("f1-score", 0),
        }
    )
    print("done.")

# =========================================================================================
# 5. TRAIN NEURAL NETWORK (batch 2)
# =========================================================================================
print("\n🧠 Training feedforward Neural Network...")

import tensorflow as tf

np.random.seed(42)
tf.random.set_seed(42)

nn_model = Sequential(
    [
        Dense(64, activation="relu", input_shape=(X_train_s.shape[1],)),
        BatchNormalization(),
        Dropout(0.3),
        Dense(32, activation="relu"),
        BatchNormalization(),
        Dropout(0.2),
        Dense(1, activation="sigmoid"),
    ]
)
nn_model.compile(optimizer="adam", loss="binary_crossentropy", metrics=["accuracy"])

history = nn_model.fit(
    X_train_s,
    y_train,
    validation_data=(X_val_s, y_val),
    epochs=20,
    batch_size=32,
    verbose=0,
)

nn_probs = nn_model.predict(X_test_s, verbose=0).flatten()
nn_preds = (nn_probs > 0.5).astype(int)
nn_key = "feedforward_nn"

model_preds[nn_key] = nn_preds
model_probs[nn_key] = nn_probs

report_dict_nn = classification_report(
    y_test,
    nn_preds,
    output_dict=True,
    target_names=le.classes_,
    zero_division=0,
)
ai_metrics_nn = report_dict_nn.get(ai_label_name, {})
human_metrics_nn = report_dict_nn.get(human_label_name, {})

results_list.append(
    {
        "Model": "Feedforward NN + BERT",
        "Accuracy": accuracy_score(y_test, nn_preds),
        "Macro F1-Score": f1_score(y_test, nn_preds, average="macro"),
        "ROC-AUC": roc_auc_score(y_test, nn_probs),
        "Precision (AI)": ai_metrics_nn.get("precision", 0),
        "Recall (AI)": ai_metrics_nn.get("recall", 0),
        "F1-Score (AI)": ai_metrics_nn.get("f1-score", 0),
        "Precision (Human)": human_metrics_nn.get("precision", 0),
        "Recall (Human)": human_metrics_nn.get("recall", 0),
        "F1-Score (Human)": human_metrics_nn.get("f1-score", 0),
    }
)

# =========================================================================================
# 6. RESULTS & BEST MODEL
# =========================================================================================
all_results_df = pd.DataFrame(results_list)
target_metric = "Macro F1-Score"

best_idx = all_results_df[target_metric].idxmax()
best_model_name = all_results_df.loc[best_idx, "Model"]
best_model_score = all_results_df.loc[best_idx, target_metric]

# Map back from pretty name to internal key
pretty_to_key = {v["Model"]: k for k, v in zip(model_preds.keys(), results_list)}
best_key_internal = None
for k, v in zip(model_preds.keys(), all_results_df["Model"]):
    if v == best_model_name:
        best_key_internal = k
        break

print("\n🏆 BEST MODEL SELECTED")
print(f"   Name : {best_model_name}")
print(f"   {target_metric}: {best_model_score:.4f}")

print("\n📋 All models (sorted):")
print(all_results_df.sort_values(target_metric, ascending=False).round(4).to_string(index=False))

# Best model detailed report
best_report = classification_report(
    y_test,
    model_preds[best_key_internal],
    output_dict=True,
    target_names=le.classes_,
)
best_model_report_df = pd.DataFrame(best_report).transpose()

# Feature importance (from Random Forest)
tree_model = models["random_forest"]
importances = tree_model.feature_importances_
feat_imp_df = (
    pd.DataFrame({"Feature": required_features, "Importance": importances})
    .sort_values("Importance", ascending=False)
    .reset_index(drop=True)
)

# Error analysis
test_subset = df_all.loc[idx_test].copy()
test_subset["True_Label"] = le.inverse_transform(y_test)
test_subset["Predicted_Label"] = le.inverse_transform(model_preds[best_key_internal])
test_subset["Prob_AI"] = model_probs[best_key_internal]

error_df = test_subset[test_subset["True_Label"] != test_subset["Predicted_Label"]]
export_cols = [orig_text_col, "True_Label", "Predicted_Label", "Prob_AI"] + required_features
error_df_clean = error_df[export_cols] if not error_df.empty else pd.DataFrame(columns=export_cols)

print(f"\n🔍 Misclassified samples: {len(error_df_clean)}")

# =========================================================================================
# 7. VISUALIZATIONS
# =========================================================================================
print("\n📈 Creating visualizations...")

# 1) Feature importance
plt.figure(figsize=(8, 5))
sns.barplot(data=feat_imp_df, x="Importance", y="Feature")
plt.title("Feature Importance (Random Forest)")
plt.tight_layout()
plt.savefig("feature_importance.png", dpi=300, bbox_inches="tight")
plt.close()

# 2) Model performance comparison
plt.figure(figsize=(12, 6))
melted_df = all_results_df.melt(
    id_vars="Model",
    value_vars=["Accuracy", "Macro F1-Score", "ROC-AUC"],
    var_name="Metric",
    value_name="Score",
)
sns.barplot(data=melted_df, x="Model", y="Score", hue="Metric")
plt.title("Model Performance Comparison")
plt.xticks(rotation=45)
plt.ylim(0, 1.05)
plt.tight_layout()
plt.savefig("model_performance.png", dpi=300, bbox_inches="tight")
plt.close()

# 3) ROC curves
plt.figure(figsize=(8, 6))
for k, probs in model_probs.items():
    fpr, tpr, _ = roc_curve(y_test, probs)
    auc_val = auc(fpr, tpr)
    label = k if len(k) <= 12 else k[:12]
    plt.plot(fpr, tpr, label=f"{label} (AUC={auc_val:.3f})")
plt.plot([0, 1], [0, 1], "k--", alpha=0.5)
plt.xlabel("False Positive Rate")
plt.ylabel("True Positive Rate")
plt.title("ROC Curves (All Models)")
plt.legend(bbox_to_anchor=(1.05, 1), loc="upper left", fontsize=8)
plt.tight_layout()
plt.savefig("roc_curves.png", dpi=300, bbox_inches="tight")
plt.close()

# 4) NN training history
history_df = pd.DataFrame(history.history)
plt.figure(figsize=(10, 4))
plt.subplot(1, 2, 1)
plt.plot(history_df["accuracy"], label="Train")
plt.plot(history_df["val_accuracy"], label="Validation")
plt.title("NN Accuracy")
plt.legend()
plt.subplot(1, 2, 2)
plt.plot(history_df["loss"], label="Train")
plt.plot(history_df["val_loss"], label="Validation")
plt.title("NN Loss")
plt.legend()
plt.tight_layout()
plt.savefig("nn_history.png", dpi=300, bbox_inches="tight")
plt.close()

print("✅ Plots saved: feature_importance.png, model_performance.png, roc_curves.png, nn_history.png")

# =========================================================================================
# 8. WORD REPORT
# =========================================================================================
print("\n📄 Generating Word report...")

doc = Document()
doc.add_heading("Arabic AI vs Human Text Classification - Phase 4", 0)

doc.add_heading("1. Data Split Summary", level=1)
table = doc.add_table(rows=len(split_summary_df) + 1, cols=len(split_summary_df.columns))
style_word_table(table)
for i, col in enumerate(split_summary_df.columns):
    table.rows[0].cells[i].text = str(col)
for i, row in split_summary_df.iterrows():
    for j, val in enumerate(row):
        table.rows[i + 1].cells[j].text = str(val)

doc.add_heading("2. Best Model", level=1)
doc.add_paragraph(f"Best model (by Macro F1-Score): {best_model_name} ({best_model_score:.4f})")

doc.add_heading("3. Model Performance Table", level=1)
all_sorted = all_results_df.sort_values(target_metric, ascending=False)
table = doc.add_table(rows=len(all_sorted) + 1, cols=len(all_sorted.columns))
style_word_table(table)
for i, col in enumerate(all_sorted.columns):
    table.rows[0].cells[i].text = str(col)
for i, row in all_sorted.iterrows():
    for j, val in enumerate(row):
        table.rows[i + 1].cells[j].text = f"{val:.4f}" if isinstance(val, (float, np.floating)) else str(val)

doc.add_heading("4. Best Model Detailed Report", level=1)
table = doc.add_table(rows=len(best_model_report_df) + 1, cols=len(best_model_report_df.columns) + 1)
style_word_table(table)
table.rows[0].cells[0].text = "Metric"
for i, col in enumerate(best_model_report_df.columns):
    table.rows[0].cells[i + 1].text = str(col)
for i, (idx, row) in enumerate(best_model_report_df.iterrows()):
    table.rows[i + 1].cells[0].text = str(idx)
    for j, val in enumerate(row):
        table.rows[i + 1].cells[j + 1].text = f"{val:.4f}" if isinstance(val, (float, np.floating)) else str(val)

doc.add_heading("5. Visualizations", level=1)
for img in ["model_performance.png", "roc_curves.png", "feature_importance.png", "nn_history.png"]:
    if os.path.exists(img):
        doc.add_picture(img, width=Inches(6))

doc.save("Final_Classification_Report.docx")
print("✅ Word report saved: Final_Classification_Report.docx")

# =========================================================================================
# 9. EXCEL REPORT
# =========================================================================================
print("\n📊 Generating Excel report...")

excel_out = "Classification_Results_Complete.xlsx"
with pd.ExcelWriter(excel_out, engine="xlsxwriter") as writer:
    split_summary_df.to_excel(writer, "01_Data_Split", index=False)
    all_results_df.sort_values(target_metric, ascending=False).round(4).to_excel(
        writer, "02_All_Results", index=False
    )
    pd.DataFrame({"Best_Model": [best_model_name], target_metric: [best_model_score]}).to_excel(
        writer, "03_Best_Model", index=False
    )
    best_model_report_df.round(4).to_excel(writer, "04_Best_Detailed_Report", index=True)
    feat_imp_df.round(4).to_excel(writer, "05_Feature_Importance", index=False)
    history_df.round(4).to_excel(writer, "06_NN_Training_History", index=True)
    if not error_df_clean.empty:
        error_df_clean.to_excel(writer, "07_Error_Analysis", index=False)
    else:
        pd.DataFrame({"Status": ["Perfect accuracy - no errors"]}).to_excel(
            writer, "07_Error_Analysis", index=False
        )

print(f"✅ Excel report saved: {excel_out}")

# =========================================================================================
# 10. SAVE MODELS + UNIFIED PREDICTOR + ZIP
# =========================================================================================
print("\n💾 Saving models and unified predictor...")

models_dict = {**models, nn_key: nn_model}
model_paths = save_all_models(models_dict, save_dir="models")
predictor_path = save_unified_predictor(
    scaler, le, model_paths, best_key_internal, save_dir="models"
)
zip_path = create_models_zip(save_dir="models", zip_name="models.zip")

# =========================================================================================
# 11. FINAL SUMMARY
# =========================================================================================
print("\n" + "=" * 80)
print("🎉 PHASE 4 COMPLETED SUCCESSFULLY")
print("=" * 80)
print(f"📊 Excel:  {excel_out}")
print("📄 Word:   Final_Classification_Report.docx")
print("📁 Models: models/  (ML + NN + scaler + label_encoder + unified_ai_detector.py)")
print(f"📦 ZIP:    {zip_path}")
print("🖼  Plots:  model_performance.png, roc_curves.png, feature_importance.png, nn_history.png")
print(f"\n🏆 Best model: {best_model_name}  |  Macro F1-Score = {best_model_score:.4f}")
print(f"🔍 Misclassified samples: {len(error_df_clean)}")
print("=" * 80)
print("Usage example:")
print("  from models.unified_ai_detector import detector")
print("  probs = detector.predict_best_proba(X_new_scaled_features)")