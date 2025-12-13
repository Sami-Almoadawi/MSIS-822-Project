import io
import os
import joblib
import pandas as pd
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from tensorflow.keras.models import Model as KerasModel

def set_table_borders(table):
    """Sets simple borders for Word tables (From Phase 1 logic)"""
    table.style = 'Table Grid'
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(2.0 if i == 0 else 1.5)

def add_df_to_word_simple(doc, df, title):
    """Adds a DataFrame to Word using simple text insertion"""
    doc.add_heading(title, level=3)
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
    set_table_borders(table)
    # Header
    for j, col_name in enumerate(df.columns):
        table.cell(0, j).text = str(col_name)
    # Body
    for i, row in df.iterrows():
        for j, value in enumerate(row):
            table.cell(i + 1, j).text = str(value)
    doc.add_paragraph('')

def add_plot_to_word(doc, fig, width_inch=6.0):
    """Saves plot to buffer and adds to Word"""
    memfile = io.BytesIO()
    fig.savefig(memfile, format='png', bbox_inches='tight', dpi=100)
    doc.add_picture(memfile, width=Inches(width_inch))
    memfile.close()
    doc.add_paragraph('')

def style_word_table(table):
    """Apply borders + blue header styling (From Phase 4 logic)"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)
    
    for border in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        edge = OxmlElement(f"w:{border}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), "000000")
        tblBorders.append(edge)

def save_all_models(models_dict, save_dir="models"):
    """Saves all ML/DL models to disk based on their type."""
    os.makedirs(save_dir, exist_ok=True)
    for model_name, model_obj in models_dict.items():
        if isinstance(model_obj, KerasModel):
            file_path = os.path.join(save_dir, f"{model_name}.keras")
            model_obj.save(file_path)
            print(f"[Saved] Keras model -> {file_path}")
        else:
            file_path = os.path.join(save_dir, f"{model_name}.pkl")
            joblib.dump(model_obj, file_path)
            print(f"[Saved] Pickle model -> {file_path}")